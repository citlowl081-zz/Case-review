"""Word document parser using python-docx."""
from parsers.base import BaseParser, ParseResult


class DocxParser(BaseParser):
    name = "docx"
    formats = ["docx", "doc"]

    def parse(self, file_path: str) -> ParseResult:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            return ParseResult(
                success=False,
                error="python-docx not installed. Run: pip install python-docx",
            )

        result = ParseResult(success=True)

        try:
            doc = DocxDocument(file_path)

            # Extract paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            result.text = "\n\n".join(text_parts)

            # Extract tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                if table_data:
                    result.tables.append(table_data)
                    # Render as text
                    lines = [" | ".join(row) for row in table_data]
                    result.tables_text.append("\n".join(lines))

        except Exception as e:
            result.success = False
            result.error = f"Word解析失败: {str(e)}"

        return result
